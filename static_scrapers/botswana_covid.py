import os
import re
import sys

import pandas as pd
from docx import Document

script_dir = os.path.abspath(os.path.dirname(sys.argv[0]) or '.')
resource_path = os.path.join(script_dir, './resources')

the_file = f'{resource_path}//BWA.Sitrep.Edited_ BWA COVID SITREP; EPI WEEK 38 Edited.docx'


def read_docx(file: str) -> {}:
    document = Document(file)
    # print(document.paragraphs)

    read_paragraphs = []
    for para in document.paragraphs:
        if len(para.text) > 0:
            read_paragraphs.append(re.sub('(?<=\d), (?=\d)', '', para.text))

    # print(read_paragraphs)

    total_second_booster = []
    total_confirmed = []
    total_deaths = []

    total_data = {}
    pat = '[\d]+[.,\d]+'
    for data in read_paragraphs:
        num = data.find("cumulative number of confirmed")
        if num != -1:
            total_confirmed.append(re.findall(pat, data[num:])[0].replace(',', '').replace(".", ""))

        num = data.find("second booster")
        if num != -1:
            total_second_booster.append(re.findall(pat, data[num:])[0].replace(',', '').replace(".", ""))

        num = data.find("Cumulative deaths")
        if num != -1:
            total_deaths.append(re.findall(pat, data[num:])[0].replace(',', '').replace(".", ""))

    total_data['total_second_booster'] = total_second_booster
    total_data['total_confirmed'] = total_confirmed
    total_data['total_deaths'] = total_deaths
    # print(total_data)

    return total_data


csv_path = os.path.join(script_dir, '../output')
df = pd.DataFrame(read_docx(the_file))
df.to_csv(f'{csv_path}//covid_total(Botswana).csv', index=False)
